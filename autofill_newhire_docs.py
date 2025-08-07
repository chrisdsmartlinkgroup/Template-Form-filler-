import os
from docx import Document
from docx.shared import Pt

def replace_placeholders(doc, data):
    """Replace all occurrences of placeholders in the Word document."""
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                # Format the replaced text
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(14)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"<<{key}>>"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            # Format the replaced text
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(14)


def fill_docx_template(input_docx_path, output_docx_path, data):
    # Load the template Word document
    doc = Document(input_docx_path)

    # Replace placeholders
    replace_placeholders(doc, data)

    # Save the updated document
    doc.save(output_docx_path)
    print(f"✅ Document filled and saved to {output_docx_path}")

if __name__ == "__main__":
    try:
        # Collect user input
        employee_name = input("Enter Employee Name: (first + last) ")
        computer_login = input("Enter Computer Login: ")
        computer_name = input("Enter Device ID: ")
        password = input("Enter Password: ")

        # Define the data for placeholders
        form_data = {
            "EMPLOYEE_NAME": employee_name,
            "COMPUTER_NAME": computer_name,
            "COMPUTER_LOGIN": computer_login,
            "PASSWORD": password
        }

        # Template paths
        template_1 = r""
        template_2 = r""

        # Output folder
        output_folder = r"C:\Users\christopher.desmond\Downloads\New Hires\New Hires Output"
        os.makedirs(output_folder, exist_ok=True)

        # Generate output file paths
        output_1 = os.path.join(output_folder, f"{employee_name.replace(' ', '_')}_first_day_form.docx")
        output_2 = os.path.join(output_folder, f"{employee_name.replace(' ', '_')}_new_hire_form.docx")

        # Fill the forms
        fill_docx_template(template_1, output_1, form_data)
        fill_docx_template(template_2, output_2, form_data)

    #error handling
    except Exception as e:
        print(f"❌ An error occurred: {e}")
