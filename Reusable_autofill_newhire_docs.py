import os
from docx import Document
from docx.shared import Pt

def replace_placeholders(doc, data):
    """Replace all occurrences of placeholders in the Word document."""
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                # Format the replaced text
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(14)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"<<{key}>>"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            # Format the replaced text
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(14)

def validate_file_path(file_path, file_type="file"):
    """Validate if a file or directory path exists."""
    if file_type == "file":
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Template file not found: {file_path}")
    elif file_type == "directory":
        if not os.path.isdir(file_path):
            raise NotADirectoryError(f"Directory not found: {file_path}")

def fill_docx_template(input_docx_path, output_docx_path, data):
    """Fill a Word document template with provided data."""
    # Load the template Word document
    doc = Document(input_docx_path)

    # Replace placeholders
    replace_placeholders(doc, data)

    # Save the updated document
    doc.save(output_docx_path)
    print(f"Document filled and saved to {output_docx_path}")

if __name__ == "__main__":
    try:
        print("=== Word Document Template Filler ===\n")
        
        # Get template file path from user
        print("Template File Setup:")
        template_path = input("Enter the full path to your Word template file: ").strip().strip('"')
        validate_file_path(template_path, "file")
        
        # Get output directory from user
        print("\nOutput Directory Setup:")
        output_folder = input("Enter the output directory path: ").strip().strip('"')
        
        # Create output directory if it doesn't exist
        try:
            os.makedirs(output_folder, exist_ok=True)
            print(f"Output directory ready: {output_folder}")
        except Exception as e:
            print(f"Could not create output directory: {e}")
            raise
        
        print("\nEmployee Information:")
        # Collect employee data
        employee_name = input("Enter Employee Name (first + last): ").strip()
        computer_login = input("Enter Computer Login (Firstname.Lastname): ").strip()
        computer_name = input("Enter Device ID: ").strip()
        password = input("Enter Password: ").strip()

        # Validate required fields
        required_fields = {
            "Employee Name": employee_name,
            "Computer Login": computer_login,
            "Device ID": computer_name,
            "Password": password
        }
        
        for field_name, field_value in required_fields.items():
            if not field_value:
                raise ValueError(f"{field_name} cannot be empty")

        # Define the data for placeholders
        form_data = {
            "EMPLOYEE_NAME": employee_name,
            "COMPUTER_NAME": computer_name,
            "COMPUTER_LOGIN": computer_login,
            "PASSWORD": password
        }

        # Generate output file path
        safe_employee_name = employee_name.replace(' ', '_').replace('/', '_').replace('\\', '_')
        output_filename = f"{safe_employee_name}_first_day_form.docx"
        output_path = os.path.join(output_folder, output_filename)

        print(f"\nProcessing template...")
        # Fill the form
        fill_docx_template(template_path, output_path, form_data)
        
        print(f"\nSuccess! Your document has been created.")
        print(f"File location: {output_path}")

    except KeyboardInterrupt:
        print("\n\nOperation cancelled by user.")
    except FileNotFoundError as e:
        print(f"\nFile Error: {e}")
        print("Make sure the template file path is correct and the file exists.")
    except NotADirectoryError as e:
        print(f"\nDirectory Error: {e}")
        print("Make sure the output directory path is correct.")
    except ValueError as e:
        print(f"\nInput Error: {e}")
        print("Please provide all required information.")
    except Exception as e:
        print(f"\nAn unexpected error occurred: {e}")
        print("Please check your file paths and try again.")
    
    input("\nPress Enter to exit...")